"""
Functions to identify compatible formatting for sequence features
which may overlap, and write the formatted sequence to a .docx
so that features can be identified.
"""

from Bio.SeqRecord import SeqRecord

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.base import EnumValue
from docx.enum.text import WD_COLOR

import logging

highlight_colours = [
    # WD_COLOR.BLUE,
    # WD_COLOR.BRIGHT_GREEN,
    WD_COLOR.GREEN,
    # WD_COLOR.PINK,
    WD_COLOR.GRAY_25,
    WD_COLOR.TEAL,
    # WD_COLOR.TURQUOISE,
    WD_COLOR.VIOLET,
    # WD_COLOR.YELLOW,
    WD_COLOR.RED,
    WD_COLOR.GRAY_50,
    WD_COLOR.DARK_YELLOW,
    WD_COLOR.DARK_RED,
    WD_COLOR.DARK_BLUE,
]
# Change font colour options so that they are all distinguishable
#  on all highlight colours
font_colours = [
    RGBColor(0xF8, 0x76, 0x6D),
    RGBColor(0xB7, 0x9F, 0x00),
    RGBColor(0x00, 0xBA, 0x38),
    RGBColor(0x00, 0xBF, 0xC4),
    RGBColor(0x61, 0x9C, 0xFF),
    RGBColor(0xF5, 0x64, 0xE3),
]
formats = [highlight_colours, font_colours, ["bold"], ["underline"], ["italic"]]
# Sorted so that formats and groups run largest to smallest
formats.sort(key=len, reverse=True)


def _make_bold(run, formatting):
    run.font.bold = True


def _make_italic(run, formatting):
    run.font.italic = True


def _make_underline(run, formatting):
    run.font.underline = True


def _colour_font(run, colour):
    run.font.color.rgb = colour


def _highlight_font(run, colour):
    run.font.highlight_color = colour


def _get_formatter(formatting):
    if isinstance(formatting, RGBColor):
        formatter = _colour_font
    elif isinstance(formatting, EnumValue):
        formatter = _highlight_font
    elif isinstance(formatting, str):
        formatters = {
            "bold": _make_bold,
            "italic": _make_italic,
            "underline": _make_underline,
        }
        formatter = formatters[formatting]
    else:
        raise ValueError(f"unknown formatting type, {formatting}")
    return formatter


def _apply_format(run, formatting):
    """Apply formatting to a specified run

    Args:
        run (docx run): The run to be formatted
        formatting (str, RGBColor, EnumValue): The formatting to be applied
            `_get_formatter()` determines how to apply the formatting based
            on its type.
    """
    formatter = _get_formatter(formatting)
    formatter(run, formatting)


def _features_per_position(seqrec, qualifier):
    """Identify all features at each position

    Args:
        seqrec (seq record): Sequence record annotated with features
        qualifier (str): The qualifier to take as a feature name, must be unique.

    Returns:
        list of lists: Each sub list reports all features overlapping
        the ith position.
    """
    position_features = []
    for i, x in enumerate(seqrec.seq):
        pf = [f.qualifiers[qualifier] for f in seqrec.features if i in f]
        position_features.append(pf)
    return position_features


def _get_feature_names(position_features):
    """Get names of each feature in order they appear.
    Names defined by `qualifier` passed to `features_per_position()`

    Args:
        position_features (list of lists): Each sublist represents all the
        features at that position in the sequence

    Returns:
        list: Each element is a feature name. They are unique and in order
    """
    fnames = [f for pos in position_features for f in pos]
    unique_fnames = []
    for f in fnames:
        if not f in unique_fnames:
            unique_fnames.append(f)
    return unique_fnames


def _check_enough_groups(adjacency, formats=formats):
    for key in adjacency.keys():
        if (len(adjacency[key]) + 1) > len(formats):
            logging.warning(
                f"{key} overlaps more features than there are formatting groups."
            )


def _get_connections(position_features):
    """Identify which features occur at the same position.
    Overlapping features must be in different groups.

    Args:
        position_features (_type_): _description_
        formats (list, optional): List of format groups to use on features.

    Returns:
        dict: each key is a feature and the items are the features
        that overlap with the key feature.
    """
    adjacency = {}
    for fname in _get_feature_names(position_features):
        # Create an empty adjacenct list for feature f
        adjacency[fname] = []
        for pf in position_features:
            # If f is at a position
            if fname in pf:
                # add those position features to f's adjacency list
                adjacency[fname].extend(pf)
                # remove f itself from f's adjacency list
                adjacency[fname].remove(fname)
        # remove duplicates
        adjacency[fname] = set(adjacency[fname])
    _check_enough_groups(adjacency)
    return adjacency


def _check_all_features_grouped(position_features, feature_groups):
    grouped_features = _get_feature_names(feature_groups)
    ufeatures = _get_feature_names(position_features)
    assert all(
        uf in grouped_features for uf in ufeatures
    ), "Not all features given a group, likely because of complex overlapping features."


def make_groups(position_features):
    """Make groups where no overlapping features share a group

    Args:
        position_features (list of lists): Each sublist represents a
        single position in the sequence and lists all the features
        that position is part of.

    Returns:
        list of lists: each sublist contains all features in a single group.
        Groups can use a single type of formatting, eg highlighting or font colour.
    """
    feature_groups = [[], [], [], [], []]
    adjacency = _get_connections(position_features)
    for fname in _get_feature_names(position_features):
        # if f already has a group skip f
        if any(fname in group for group in feature_groups):
            continue
        for group in feature_groups:
            group
            if all(f_connection not in group for f_connection in adjacency[fname]):
                group.append(fname)
                break
    _check_all_features_grouped(position_features, feature_groups)
    return feature_groups


def _check_enough_formats(group, group_formats):
    assert len(group) <= len(
        group_formats
    ), f"{group} contains more features than there are available formats"


def _make_feature_formats(feature_groups, formats):
    """Assign formatting to each feature

    Args:
        feature_groups (list): list of groups
        formats (list of lists): each sublist contains all available values of a
        type of formatting, eg font colours.

    Returns:
        dict: Keys are features, items are their format.
    """
    feature_formats = {}
    for i, group in enumerate(feature_groups):
        _check_enough_formats(group, formats[i])
        for j, feature in enumerate(group):
            feature_formats[feature] = formats[i][j]
    return feature_formats


def _write_position_features(position_features, seqrec, feature_formats, p):
    """Write a sequence formatted according to its features. Note that
    all sequences are written as upper case.

    Args:
        position_features (list of lists): Each sublist represents a
        single position in the sequence and lists all the features
        that position is part of.
        seqrec (SeqRecord): The SeqRecord that `position_features` describe.
        feature_formats dict: Keys are features, items are their format.
        p (docx paragraph): The docx paragraph element to write to
    """
    for i, pfs in enumerate(position_features):
        run = p.add_run(seqrec[i].upper())
        run.font.size = Pt(8)
        for pf in pfs:
            _apply_format(run, feature_formats[pf])


def _write_heading(seqrec, doc):
    """Generate a small heading with sequence name and description

    Args:
        seqrec (SeqRecord): The sequence record
        doc (docx document): The document element to write to
    """
    doc.add_heading(seqrec.id, level=1)
    p = doc.add_paragraph(seqrec.name)
    p.add_run(f"\n{seqrec.description}")


def _write_legend(feature_formats, doc):
    """Write a legend of features and their formats"""
    p = doc.add_paragraph("Legend")
    for feature in feature_formats.keys():
        run = p.add_run(f"\n{feature}")
        _apply_format(run, feature_formats[feature])


def write_crazyseq(seqrec, qualifier, doc, formats=formats):
    """Write a single sequence with annotated features to an existing
    document element. If you want crazydocs to handle document
    creation use `write_crazydoc()` even for single sequences.

    Args:
        seqrec (SeqRecord): The sequence to write, with features.
        qualifier (str): The key of the .qualifier to use as feature names.
        doc (docx document): The docx document element to write to.
        formats (list, optional): List of format groups to use on features.
    """
    _write_heading(seqrec, doc)
    # Determine feature groups and formats
    position_features = _features_per_position(seqrec, qualifier)
    feature_groups = make_groups(position_features)
    feature_groups.sort(key=len, reverse=True)
    feature_formats = _make_feature_formats(feature_groups, formats)
    p = doc.add_paragraph()
    _write_position_features(position_features, seqrec, feature_formats, p)
    _write_legend(feature_formats, doc)


def write_crazydoc(seqs, qualifier, path, formats=formats):
    """Write one or more SeqRecords annotated with features to a new
    document.
    WILL OVERWRITE EXISTING DOUCMENTS WITHOUT WARNING.

    Args:
        seqs (list): List of SeqRecords to write out
        qualifier (str): The key to use for feature names from `.qualifiers`
        path (str): The path to write the document to. Will overwrite existing documents!
        formats (list, optional): List of format groups to use on features.

    Examples:
        A custom list of docx formats can be supplied to the `formats` argument.
        If `formats` is not declared a default palette is used.

        >>>from docx.shared import RGBColor
        >>>from docx.enum.text import WD_COLOR
        >>>custom_formats = [
                [
                    WD_COLOR.PINK,
                    WD_COLOR.TEAL,
                    WD_COLOR.YELLOW,
                    WD_COLOR.GRAY_50,
                    WD_COLOR.VIOLET,
                    WD_COLOR.TURQUOISE,
                ],
                [
                    RGBColor(0xFF, 0x00, 0x00),
                    RGBColor(0x00, 0xFF, 0x00),
                    RGBColor(0x00, 0x00, 0xFF),
                ],
                ["bold"],
                ["underline"],
                ["italic"],
            ]
        >>>write_crazydoc(seq, 'product', 'custom_colours.docx', formats=custom_formats)

    """
    doc = Document()
    if isinstance(seqs, SeqRecord):
        seqs = [seqs]
    for record in seqs:
        write_crazyseq(record, qualifier, doc, formats)
    doc.save(path)


def _test_colours(path, highlight_colours=highlight_colours, font_colours=font_colours):
    """Test all colour combinations

    Args:
        path (str): Path to save colour test docx
        highlight_colours (list of WC_COLOR_INDEX, optional): List of used highlight
            colours in docx WD_COLOR_INDEX form. Defaults to highlight_colours.
        font_colours (list of RGBColours, optional): List of used font colours in docx
            RGBColor form. Defaults to font_colours.
    """
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run(f"\nlorem ipsum")
    for font_col in font_colours:
        run = p.add_run(f"\nlorem ipsum")
        _apply_format(run, font_col)
    for highlight_col in highlight_colours:
        p = doc.add_paragraph()
        run = p.add_run(f"\nlorem ipsum")
        _apply_format(run, highlight_col)
        for font_col in font_colours:
            run = p.add_run(f"\nlorem ipsum")
            _apply_format(run, highlight_col)
            _apply_format(run, font_col)
    doc.save(path)
